# Standard library imports
import os
import argparse
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed

# Third-party library imports
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from tqdm import tqdm
from dotenv import dotenv_values
from openai import OpenAI
from tenacity import retry, wait_exponential, stop_after_attempt

# 配置日志
logging.basicConfig(
    filename='process.log',
    filemode='a',
    format='%(asctime)s - %(levelname)s - %(message)s',
    level=logging.INFO
)

def load_prompt(prompt_file):
    """
    加载指定的提示语文件，并返回提示语字符串。
    """
    if not os.path.exists(prompt_file):
        logging.error(f"提示文件 {prompt_file} 不存在。")
        raise FileNotFoundError(f"提示文件 {prompt_file} 不存在。")
    with open(prompt_file, 'r', encoding='utf-8') as file:
        prompt = file.read()
    logging.info(f"成功加载提示文件: {prompt_file}")
    return prompt

def initialize_openai_client():
    """
    初始化 OpenAI 客户端，并返回客户端实例以及API调用参数。
    """
    config = dotenv_values(".env")
    api_key = config.get("OPENAI_API_KEY")
    if not api_key:
        logging.error("未找到 OPENAI_API_KEY，请在 .env 文件中配置。")
        raise ValueError("未找到 OPENAI_API_KEY，请在 .env 文件中配置。")
    
    # 从配置中获取API调用参数，提供默认值
    temperature = float(config.get("TEMPERATURE", 0.3))
    max_tokens = int(config.get("MAX_TOKENS", 1024))
    top_p = float(config.get("TOP_P", 1))
    frequency_penalty = float(config.get("FREQUENCY_PENALTY", 0))
    presence_penalty = float(config.get("PRESENCE_PENALTY", 0))
    model = config.get("MODEL", "gpt-4o-2024-08-06")
    
    client = OpenAI(
        api_key=api_key,
    )
    
    api_params = {
        "temperature": temperature,
        "max_tokens": max_tokens,
        "top_p": top_p,
        "frequency_penalty": frequency_penalty,
        "presence_penalty": presence_penalty,
        "model": model
    }
    
    logging.info(f"初始化 OpenAI 客户端成功，使用模型: {model}")
    return client, api_params

@retry(wait=wait_exponential(min=5, max=60), stop=stop_after_attempt(3))
def review_content(content, init_prompt, client, api_params):
    """
    使用 OpenAI API 评审内容，并返回评审结果。
    """
    try:
        response = client.chat.completions.create(
            model=api_params["model"],
            messages=[
                {
                    "role": "system",
                    "content": init_prompt
                },
                {
                    "role": "user",
                    "content": f"```\n{content}\n```"
                }
            ],
            temperature=api_params["temperature"],
            max_tokens=api_params["max_tokens"],
            top_p=api_params["top_p"],
            frequency_penalty=api_params["frequency_penalty"],
            presence_penalty=api_params["presence_penalty"]
        )
        result = response.choices[0].message.content.strip()
        logging.info("成功调用 OpenAI API 评审内容。")
        return result
    except Exception as e:
        logging.error(f"OpenAI API 调用失败: {e}")
        raise  # 重新抛出异常以触发重试

def process_file(file, init_prompt, client, api_params, prompt_name, output_dir):
    """
    处理单个 .docx 文件，并生成评审结果文档。
    """
    try:
        # 打开原始文档
        doc = Document(file)

        # 提取文档内容
        content = "\n".join([para.text for para in doc.paragraphs])

        # 调用 OpenAI API 进行评审
        reviewed_content = review_content(content, init_prompt, client, api_params)

        if reviewed_content:
            # 创建一个新文档
            new_doc = Document()

            # 设置默认样式为 SimSun（宋体）
            style = new_doc.styles['Normal']
            style.font.name = 'SimSun'
            style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            style.font.size = Pt(12)
            style.font.color.rgb = RGBColor(0, 0, 0)

            # 添加学生原文
            new_doc.add_heading('学生原文', level=1)
            for para in doc.paragraphs:
                new_doc.add_paragraph(para.text)

            # 添加分隔符
            new_doc.add_page_break()
            new_doc.add_heading('评审结果', level=1)

            # 添加评审内容
            for line in reviewed_content.split('\n'):
                new_doc.add_paragraph(line)

            # 获取原文件名（不含扩展名）
            original_base_name = os.path.splitext(os.path.basename(file))[0]

            # 生成新的文件名，直接连接原文件名和提示语文件名（不含下划线）
            new_filename = f"{original_base_name}{prompt_name}.docx"
            new_filepath = os.path.join(output_dir, new_filename)

            # 保存新文档
            new_doc.save(new_filepath)
            logging.info(f"已生成评审文档: {new_filepath}")
            return new_filepath
        else:
            logging.warning(f"由于 API 错误，跳过文件 {file}。")
            return None
    except Exception as e:
        logging.error(f"处理文件 {file} 时出错：{e}")
        return None

def process_documents_parallel(init_prompt, client, api_params, prompt_name, docs_path, output_dir, max_workers=5):
    """
    并行处理指定目录下的所有 .docx 文件，并生成评审结果文档。
    """
    # 获取未处理的 .docx 文件列表
    doc_files = [os.path.join(docs_path, f) for f in os.listdir(docs_path) if f.endswith('.docx')]
    
    if not doc_files:
        logging.info(f"指定目录 '{docs_path}' 下没有需要处理的 .docx 文件。")
        return

    failed_files = []

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        future_to_file = {executor.submit(process_file, file, init_prompt, client, api_params, prompt_name, output_dir): file for file in doc_files}
        for future in tqdm(as_completed(future_to_file), total=len(doc_files), desc="Processing documents"):
            file = future_to_file[future]
            try:
                result = future.result()
                if result:
                    logging.info(f"成功处理文件: {file}")
                else:
                    failed_files.append(file)
            except Exception as e:
                logging.error(f"处理文件 {file} 时出现异常：{e}")
                failed_files.append(file)
    
    if failed_files:
        logging.warning("以下文件处理失败，可以稍后重试：")
        for file in failed_files:
            logging.warning(f"- {file}")

def main():
    # 设置命令行参数解析
    parser = argparse.ArgumentParser(description="评审学生作文的脚本")
    parser.add_argument(
        '--prompt',
        type=str,
        required=True,
        help="提示语文件的路径（.txt 文件）"
    )
    parser.add_argument(
        '--docs',
        type=str,
        required=True,
        help="要遍历的文档目录路径"
    )
    parser.add_argument(
        '--model',
        type=str,
        default=None,
        help="使用的 OpenAI 模型名称（可选，优先级低于 .env 文件中的 MODEL 配置）"
    )
    parser.add_argument(
        '--max_workers',
        type=int,
        default=5,
        help="并行处理的最大工作线程数"
    )
    args = parser.parse_args()

    # 提取提示语文件的名称（不含扩展名）
    prompt_base_name = os.path.splitext(os.path.basename(args.prompt))[0]

    # 初始化 OpenAI 客户端和API调用参数
    try:
        client, api_params = initialize_openai_client()
        # 如果命令行指定了模型，则覆盖配置文件中的模型
        if args.model:
            api_params["model"] = args.model
            logging.info(f"使用命令行指定的模型: {args.model}")
    except Exception as e:
        logging.error(f"初始化 OpenAI 客户端失败: {e}")
        return

    # 加载提示语
    try:
        init_prompt = load_prompt(args.prompt)
    except Exception as e:
        logging.error(f"加载提示语文件失败: {e}")
        return

    # 获取文档目录路径
    docs_path = args.docs
    if not os.path.exists(docs_path) or not os.path.isdir(docs_path):
        logging.error(f"指定的文档目录路径 '{docs_path}' 不存在或不是一个目录。")
        return

    # 创建输出目录 'Completed'（如果不存在）
    output_dir = os.path.join(docs_path, 'Completed')
    if not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir)
            logging.info(f"创建输出目录: {output_dir}")
        except Exception as e:
            logging.error(f"创建输出目录 {output_dir} 失败: {e}")
            return

    # 并行处理文档
    process_documents_parallel(init_prompt, client, api_params, prompt_base_name, docs_path, output_dir, args.max_workers)

if __name__ == "__main__":
    main()
